import sys
import re
import os
from docx import Document
import firebase_admin
from firebase_admin import credentials, firestore
"""
yearbook_parser.py

קורא קובץ Yearbook (DOCX) ומייבא אותו ל-Firestore:
- מזהה סמסטרים
- שומר קורסים (קוד, שם, שעות, נ"ז)
- יוצר קשרי קורסים: קדם / צמוד
- משלים שמות קורסים חסרים בקשרי קדם

שימוש:
python yearbook_parser.py <docx_path> <yearbook_id> <yearbook_label>
"""
# ==============================
# Firebase init (ENV based)
# ==============================
if not firebase_admin._apps:
    cred = credentials.Certificate({
        "type": "service_account",
        "project_id": os.environ["FIREBASE_PROJECT_ID"],
        "client_email": os.environ["FIREBASE_CLIENT_EMAIL"],
        "private_key": os.environ["FIREBASE_PRIVATE_KEY"].replace("\\n", "\n"),
        "token_uri": "https://oauth2.googleapis.com/token",
    })
    firebase_admin.initialize_app(cred)

db = firestore.client()

# ==============================
# Helpers
# ==============================
def normalize(s):
    if not s:
        return ""
    return re.sub(r"\s+", " ", s.replace("\u00A0", " ")).strip()

def is_course_code(text):
    return re.fullmatch(r"\d{5,6}", text or "") is not None

def safe_cell(row, idx, dash_as_zero=False):
    if idx is None or idx >= len(row.cells):
        return None

    v = normalize(row.cells[idx].text)
    if v == "":
        return None
    if v == "-":
        return 0 if dash_as_zero else None

    if re.fullmatch(r"\d+(\.\d+)?", v):
        return float(v) if "." in v else int(v)

    return None

def safe_hours(value):
    if isinstance(value, (int, float)) and 0 <= value <= 10:
        return value
    return None

def find_col(headers, *needles):
    for i, h in enumerate(headers):
        for n in needles:
            if n in h:
                return i
    return None

def find_relation_col(headers):
    for i, h in enumerate(headers):
        if "קדם" in h or "צמוד" in h:
            return i
    return None

# ==============================
# Relations (underline support)
# ==============================
def extract_relations_from_docx_cell(cell, course_name_map):
    relations = []

    for p in cell.paragraphs:
        line_text = normalize(p.text)
        if not line_text:
            continue

        codes = re.findall(r"\b\d{5,6}\b", line_text)
        if not codes:
            continue

        is_underlined = False
        for run in p.runs:
            if run.font.underline:
                is_underlined = True
                break

        rel_type = "COREQUISITE" if is_underlined else "PREREQUISITE"

        for code in codes:
            relations.append({
                "courseCode": code,
                "courseName": course_name_map.get(code),
                "type": rel_type
            })

    uniq = {}
    for r in relations:
        uniq[r["courseCode"]] = r

    return list(uniq.values())

# ==============================
# Core logic
# ==============================
def process_docx(doc, yearbook_id, yearbook_label):
    table_map = {t._element: t for t in doc.tables}

    root = db.collection("yearbooks").document(yearbook_id)
    root.set(
        {
            "yearbookId": yearbook_id,
            "displayName": yearbook_label,
        },
        merge=True,
    )

    required = root.collection("requiredCourses")

    current_sem = None
    created_semesters = set()
    course_name_map = {}
    pending_relations = []

    for block in doc.element.body:
        # --- Detect semester ---
        if block.tag.endswith("p"):
            text = normalize("".join(t.text for t in block.xpath(".//w:t")))
            m = re.search(r"סמסטר\s*([1-8])", text)
            if m:
                current_sem = int(m.group(1))
                if current_sem not in created_semesters:
                    required.document(f"semester_{current_sem}").set(
                        {"semesterNumber": current_sem},
                        merge=True,
                    )
                    created_semesters.add(current_sem)

        # --- Courses table ---
        elif block.tag.endswith("tbl") and current_sem:
            table = table_map.get(block)
            if not table or not table.rows:
                continue

            headers = [normalize(c.text) for c in table.rows[0].cells]
            if "שם הקורס" not in " ".join(headers):
                continue

            code_i, name_i = 0, 1

            # ✅ זיהוי נכון של עמודות
            lec_i  = find_col(headers, "הרצאה", "ה")
            prac_i = find_col(headers, "תרגול", "ת")
            lab_i  = find_col(headers, "מעבדה", "מ")
            cred_i = find_col(headers, 'נ"ז', "נקודות", "נ")
            rel_i  = find_relation_col(headers)

            for row in table.rows[1:]:
                code = normalize(row.cells[code_i].text)
                name = normalize(row.cells[name_i].text)

                if not is_course_code(code):
                    continue

                if name:
                    course_name_map[code] = name

                course_ref = (
                    required.document(f"semester_{current_sem}")
                    .collection("courses")
                    .document(code)
                )

                lecture_hours  = safe_hours(safe_cell(row, lec_i, True))
                practice_hours = safe_hours(safe_cell(row, prac_i, True))
                lab_hours      = safe_hours(safe_cell(row, lab_i, True))

                course_ref.set(
                    {
                        "courseCode": code,
                        "courseName": name,
                        "lectureHours": lecture_hours,
                        "practiceHours": practice_hours,
                        "labHours": lab_hours,
                        "credits": safe_cell(row, cred_i),
                    },
                    merge=True,
                )

                # --- Relations ---
                if rel_i is not None:
                    relations = extract_relations_from_docx_cell(
                        row.cells[rel_i],
                        course_name_map
                    )

                    for r in relations:
                        rel_ref = course_ref.collection("relations").document(
                            r["courseCode"]
                        )
                        rel_ref.set(r, merge=True)

                        if r.get("courseName") is None:
                            pending_relations.append(
                                (rel_ref, r["courseCode"])
                            )

    # --- Resolve pending relation names ---
    if pending_relations:
        batch = db.batch()
        for ref, code in pending_relations:
            if code in course_name_map:
                batch.set(ref, {"courseName": course_name_map[code]}, merge=True)
        batch.commit()

# ==============================
# Entry point
# ==============================
if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: yearbook_parser.py <docx_path> <yearbook_id> <yearbook_label>")
        sys.exit(1)

    file_path = sys.argv[1]
    yearbook_id = sys.argv[2]
    yearbook_label = sys.argv[3]

    doc = Document(file_path)
    process_docx(doc, yearbook_id, yearbook_label)

    print("OK")
